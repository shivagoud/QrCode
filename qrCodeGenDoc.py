import Tkinter
from Tkinter import *
import Tkconstants
import tkFileDialog
import sys
import pyqrcode
import png
import xlrd
from xlrd import xldate
from xlrd.sheet import ctype_text
from threading import Thread
import docx
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

class Gui:
    def __init__(self):
        self.top = Tkinter.Tk()
        self.top.iconbitmap('qrcode.ico')
    #save document    
    def beToFe(self,document):
        docName = tkFileDialog.asksaveasfilename(defaultextension='.docx',
                                        filetypes=[("MS Doc Files","*.docx;*.doc")],
                                        title="File Browser")
        try :
            document.save(docName)
            self.displayTextOnGui("Saved the document to "+docName)
        except Exception, e:
            self.displayTextOnGui("Ouch..!!I can't save the file.!!",'red')
            self.displayTextOnGui("Cannot create/edit the given Filename.!!",'red')
            
    #display on the Top Gui
    def displayTextOnGui(self,string,fg='black') :
        L1 = Label(self.F1,text="+++++++"+string,fg=fg).pack(side=TOP,anchor='w')
        self.F1.pack(side=TOP,anchor='w',fill=BOTH)
        #self.canvas.yview("scroll",1,"units")
        self.canvas.yview_moveto(1.0)
        self.top.update()
    #display the chosen one
    def displayFileName(self,fileName):
        self.displayTextOnGui("Path to the chosen one : "+fileName)
        return

        
    #choose a file
    def getFileName(self):
        return tkFileDialog.askopenfilename(defaultextension='.xlsx',
                                        filetypes=[("Excel Files","*.xlsx;*.xls")],
                                        title="File Browser")
    def checkAll(self,var1,sheet_vars):
        if var1.get():
            for shname, c in sheet_vars.iteritems():
                c.set(1)
    #Choose the file and handle it.
    def handleFile(self):
        w = self.getFileName()
        if w:
            filename = w;
            self.displayFileName(filename)
            #get sheet names
            try :
                wb = xlrd.open_workbook(filename)
            except Exception, e:
                guiObj.displayTextOnGui(inputfile+" is not a proper excel file")
                exit(1)
            self.sheet_vars={}
            self.check_list=[]
            L3 = Label(self.F1,text="Select The Sheets To Generate QR:")
            L3.pack(side=TOP,anchor='w',pady=10)
            var1 = IntVar()
            cAll = Checkbutton(self.F1, text="Select All", variable=var1, command = lambda:self.checkAll(var1,self.sheet_vars))
            cAll.pack(side=TOP,anchor='w',padx=20)
            for i in range(0,wb.nsheets-1):
                sh = wb.sheet_by_index(i)
                shname = sh.name
                self.sheet_vars[shname] = IntVar()
                c = Checkbutton(self.F1, text =shname+" ("+str(sh.nrows)+")", variable = self.sheet_vars[shname])
                c.pack(side=TOP,anchor='w',padx=40)
                self.check_list.append(c)
            B2 = Button(self.F1, text="Generate QR",command=lambda: QrCodeGen().lockAndLoad(wb,self))
            B2.pack(side=TOP,anchor='w',padx=30,pady=5)
            self.F1.pack(side=TOP,fill=X)
            self.top.update()
        return
    def onFrameConfigure(self,event):
        self.canvas.configure(scrollregion=self.canvas.bbox("all"))
    def buildGui(self):
        #set the name for top window
        self.top.wm_title("QR Code Generator")
        self.top.geometry("600x700")
        topF = Frame(self.top,width=600,height=700)
        topF.pack(side=TOP,fill=BOTH)
        self.canvas = canvas = Canvas(topF,borderwidth=0,width=600,height=700)
        midF = Frame(canvas)
        vsb = Scrollbar(topF, orient="vertical", command=canvas.yview)
        canvas.configure(yscrollcommand=vsb.set)
        vsb.pack(side="right", fill="y")
        canvas.pack(side="left", fill=BOTH, expand=True)
        canvas.create_window((4,4), window=midF, anchor="nw", 
                                  tags="self.frame")
        midF.bind("<Configure>", self.onFrameConfigure)
        #add Frame for path entry widget
        self.F1 = Frame(midF)
        #frame for button widget
        F2 = Frame(midF)
        L2 = Label(F2, text="Choose a file to generate QR code: ")
        #button because i like clicking;)
        B1 = Button(F2, text="Get The File!!",command=lambda: self.handleFile())
        L2.pack(side=LEFT,anchor='w')
        B1.pack(side=LEFT,anchor='e')
        #offset for incomplete label sheet
        F3 = Frame(midF)
        L3 = Label(F3,text="Offset For The First Label : ")
        lst1 = ['0','1','2','3','4','5','6','7','8','9','10','11']
        self.varOffset = StringVar()
        self.varOffset.set('0')
        drop = OptionMenu(F3,self.varOffset,*lst1)
        L3.pack(side=LEFT)
        drop.pack(side=LEFT)
        F3.pack(side=TOP,fill=X)
        F2.pack(side=TOP,fill=X,pady=10)
        #get set go!!
        self.top.mainloop()




class QrCodeGen():
    #core funtion to generate QR Code
    def lockAndLoad(self,wb,guiObj):

        document = Document("template")
        p = document.paragraphs[0]._element
        p.getparent().remove(p)
        p._p = p._element = None
        section = document.sections[0]
        section.top_margin = Inches(0.18)
        section.bottom_margin = 0
        table = document.add_table(1,2)
        table.style = "qrstyle"
        table.allow_autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        celli = 0
        rowCount = 1
        offset = int(guiObj.varOffset.get())
        while (offset > 0) :
            if celli == 2:
                row = table.add_row()
                rowCount+=1
                celli = 0
            row = table.rows[rowCount-1]
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), "45mm")
            trHeight.set(qn('w:hRule'), "atLeast")
            trPr.append(trHeight)
            celli += 1
            offset -= 1
        for i in range(0,wb.nsheets-1):
            sh = wb.sheet_by_index(i)
            shname = sh.name
            var = guiObj.sheet_vars[shname]
            if var.get() == 0:
                continue
            #addHeader = 0 #no header for now
            guiObj.displayTextOnGui("Processing sheet : "+shname)
            accCol = 10 #default hard code

            for idx, cell_obj in enumerate(sh.row(0)):
                if "accession" in str(sh.cell(0,idx).value).lower():
                        accCol = idx
                        break
            procRows=0
            for j in range(1,sh.nrows):
                if(sh.nrows == 1):
                    break
                row = sh.row(j)
                qr_string=''
                for idx, cell_obj in enumerate(row):
                    if cell_obj.ctype in (xlrd.XL_CELL_EMPTY, xlrd.XL_CELL_BLANK):
                        continue
                    if(cell_obj.ctype == xlrd.XL_CELL_DATE):
                        cell_obj.value = xldate.xldate_as_datetime(cell_obj.value,0).date().strftime('%Y-%b-%d')
                    if(cell_obj.ctype == xlrd.XL_CELL_NUMBER):
                        cell_obj.value = int(cell_obj.value)
                    qr_string += str(sh.cell(0,idx).value)+ " : " +str(cell_obj.value)+"\n"
                if qr_string == '':
                    continue
                if str(sh.cell(j,0).value) == '':
                    continue
                #if addHeader == 1 :
                    #head = document.add_heading(shname)
                    #head.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    #addHeader = 0
                big_code = pyqrcode.create(qr_string, error='L', version=11, mode='binary')
                big_code.png('code.png', scale=1, module_color=[0, 0, 0, 128], background=[0xff, 0xff, 0xff])
                if celli == 2:
                    row = table.add_row()
                    rowCount+=1
                    celli = 0
                row = table.rows[rowCount-1]
                tr = row._tr
                trPr = tr.get_or_add_trPr()
                trHeight = OxmlElement('w:trHeight')
                trHeight.set(qn('w:val'), "45mm")
                trHeight.set(qn('w:hRule'), "atLeast")
                trPr.append(trHeight)
                cell = row.cells[celli]
                cell.width = Inches(4.02)
                para = cell.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                font = run.font
                font.size = Pt(9)
                run.add_picture("code.png")
                run.add_break()
                sname = str("Title : "+sh.cell(j,0).value)
                accNoStr = str("Accession No : "+str(int(sh.cell(j,accCol).value)))
                run.add_text(sname)
                run.add_break()
                run.add_text(accNoStr)
                #run.add_break()
                celli += 1
                procRows+=1
                #dirty hack to keep gui updated
                guiObj.top.update()
            guiObj.displayTextOnGui("Processed "+str(procRows)+" Entries")
        guiObj.beToFe(document)
        
            
Gui().buildGui()
     
